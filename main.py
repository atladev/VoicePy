import os
import time
from pathlib import Path
from docx import Document
import torch
from shutil import copy2, rmtree
from TTS.api import TTS
from TTS.utils.synthesizer import Synthesizer
import io
from contextlib import redirect_stdout
import streamlit as st
import psutil
import re

# =============================
# CONSOLE STATUS / LOG HELPERS
# =============================
_last_standby_print = 0.0

def _console(status: str, msg: str = ""):
    ts = time.strftime('%Y-%m-%d %H:%M:%S')
    line = f"[{ts}] [{status.upper()}] {msg}"
    try:
        print(line, flush=True)
    except Exception:
        pass

def _console_standby_throttled(msg: str, interval: float = 8.0):
    global _last_standby_print
    now = time.time()
    if now - _last_standby_print >= interval:
        _console("STANDBY", msg)
        _last_standby_print = now

# =============================
# GENERAL SETTINGS
# =============================
p = psutil.Process(os.getpid())
try:
    p.nice(psutil.HIGH_PRIORITY_CLASS)
except Exception:
    pass

os.environ["COQUI_TOS_AGREED"] = "1"

DEFAULT_MODEL = "tts_models/multilingual/multi-dataset/xtts_v2"
DEFAULT_DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

# Default paths - adjust these to your environment
DOWNLOAD_PATH = "./output_audios"
VOICE_PATH = "./voices"
error_texts = []

# =============================
# GLOBAL LOCK SYSTEM
# =============================
LOCK_FILE = Path("app_in_use.lock")

def set_app_status(in_use: bool):
    """Creates or removes lock file to signal usage."""
    if in_use:
        LOCK_FILE.write_text("IN_USE")
    else:
        if LOCK_FILE.exists():
            LOCK_FILE.unlink(missing_ok=True)

def is_app_in_use() -> bool:
    """Checks if the system is in use by another user."""
    return LOCK_FILE.exists()

# =============================
# UTILITIES
# =============================
def sanitize_name(name: str) -> str:
    base = re.sub(r"[\\/:*?\"<>|]", "_", name)
    base = re.sub(r"\s+", " ", base).strip()
    return base

def list_wav_files(folder: str):
    try:
        p = Path(folder)
        if not p.exists():
            return []
        return [str(f.name) for f in p.glob("*.wav")]
    except Exception:
        return []

params = {
    "remove_trailing_dots": True,
    "voice": "",
    "language": "pt",
    "model_name": DEFAULT_MODEL,
    "device": DEFAULT_DEVICE,
}

def new_split_into_sentences(self, text):
    sentences = self.seg.segment(text)
    if params['remove_trailing_dots']:
        sentences_without_dots = []
        for sentence in sentences:
            if sentence.endswith('.') and not sentence.endswith('...'):
                sentence = sentence[:-1]
            sentences_without_dots.append(sentence)
        return sentences_without_dots
    else:
        return sentences

Synthesizer.split_into_sentences = new_split_into_sentences

# =============================
# MODEL LOADING
# =============================
@st.cache_resource(show_spinner=True)
def load_model(model_name: str, device: str):
    return TTS(model_name).to(device)

def flush_tts_cache():
    """Forces unloading of current model from cache and GPU."""
    try:
        _console("FLUSH", "Clearing TTS model cache...")
        st.cache_resource.clear()
        torch.cuda.empty_cache()
    except Exception as e:
        _console("ERROR", f"Failed to clear cache: {e}")

# =============================
# DOCX READING
# =============================
@st.cache_data(show_spinner=False)
def load_text(file_path):
    doc = Document(file_path)
    return [paragraph.text.replace('.', ',.') for paragraph in doc.paragraphs if paragraph.text.strip()]

def generate_audio_filename(index):
    return f"audio_{index + 1}.wav"

# =============================
# GENERATION WITH LOGGING
# =============================
def tts_to_file_logged(model, text: str, out_path: str, language: str, speaker_wav: str, speed: float = 0.85):
    output_buffer = io.StringIO()
    with redirect_stdout(output_buffer):
        model.tts_to_file(
            text=text,
            file_path=out_path,
            speaker_wav=speaker_wav,
            language=language,
            speed=speed,
        )
    return output_buffer.getvalue()

# =============================
# STREAMLIT APP
# =============================
def main_app():
    _console_standby_throttled("Application started...")
    st.set_page_config(page_title="TTS Audio Generator - Free Credits: 2000", layout="wide")

    # Clean up old lock file on initialization if needed
    if is_app_in_use():
        try:
            mod_time = time.time() - LOCK_FILE.stat().st_mtime
            if mod_time > 600:  # More than 10 minutes without use? Clean it up.
                LOCK_FILE.unlink(missing_ok=True)
                _console("INFO", "Old lock file removed automatically.")
        except Exception:
            pass

    st.title("TTS Audio Generator - Free Credits: 2000")
    st.caption("version 1.04")

    # ===================================
    # ACTIVE USE WARNING (GLOBAL BANNER)
    # ===================================
    app_status_placeholder = st.empty()
    # Always starts as available
    app_status_placeholder.markdown(
        "<div style='background-color:#4CAF50;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>✅ Available! No one is using it right now.</div>",
        unsafe_allow_html=True
    )

    with st.sidebar:
        st.subheader("Settings")

        language = st.radio("Narration Language", options=["pt", "es", "en"], index=0, horizontal=True)
        params["language"] = language

        voice_dir = st.text_input("Voice folder path (.wav files)", value=VOICE_PATH)
        wavs = list_wav_files(voice_dir)

        if not wavs:
            st.warning("No .wav files found in this folder.")
        else:
            selected_wav_name = st.selectbox("Choose voice (.wav file)", wavs)
            selected_wav_path = str(Path(voice_dir) / selected_wav_name)

            # If voice changed, force model flush
            if params.get("voice") and params["voice"] != selected_wav_path:
                flush_tts_cache()

            params["voice"] = selected_wav_path

            st.markdown("**Preview selected voice**")
            try:
                with open(selected_wav_path, "rb") as f:
                    st.audio(f.read(), format="audio/wav")
            except Exception as e:
                st.error(f"Error loading voice preview: {e}")

        with st.expander("Advanced options"):
            params["model_name"] = st.text_input("Coqui Model", value=params["model_name"])
            params["device"] = st.selectbox("Device", ["cuda", "cpu"], index=0 if DEFAULT_DEVICE == "cuda" else 1)
            params["remove_trailing_dots"] = st.checkbox("Remove simple period from end of sentences", value=True)

    st.divider()

    st.markdown("### Quick test of selected voice (TTS)")
    sample_text = st.text_input("Test text", value="This is a narration test.")
    if st.button("Play generated sample"):
        if not params.get("voice"):
            st.error("Please select a voice first.")
        elif is_app_in_use():
            st.error("Another user is using the app now. Wait for the green indicator.")
        else:
            set_app_status(True)
            app_status_placeholder.markdown(
                "<div style='background-color:#ff4d4d;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>🚫 In use! Don't touch anything right now.</div>",
                unsafe_allow_html=True
            )
            try:
                with st.spinner("Generating sample..."):
                    model = load_model(params["model_name"], params["device"])
                    tmp_sample = Path("./_tmp_sample.wav")
                    _ = tts_to_file_logged(
                        model,
                        text=sample_text,
                        out_path=str(tmp_sample),
                        language=params["language"],
                        speaker_wav=params["voice"],
                        speed=0.9,
                    )
                    with open(tmp_sample, "rb") as f:
                        st.audio(f.read(), format="audio/wav")
            finally:
                if tmp_sample.exists():
                    tmp_sample.unlink(missing_ok=True)
                set_app_status(False)
                app_status_placeholder.markdown(
                    "<div style='background-color:#4CAF50;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>✅ Available! No one is using it right now.</div>",
                    unsafe_allow_html=True
                )

    st.divider()

    uploaded_file = st.file_uploader("Select .docx file", type="docx")
    if uploaded_file is not None:
        if is_app_in_use():
            st.error("Another user is generating audio now. Wait for the green indicator.")
            return

        set_app_status(True)
        app_status_placeholder.markdown(
            "<div style='background-color:#ff4d4d;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>🚫 In use! Don't touch anything right now.</div>",
            unsafe_allow_html=True
        )

        try:
            _console("EXECUTING", f"File received: {uploaded_file.name}")
            original_name = sanitize_name(uploaded_file.name)
            base_name = os.path.splitext(original_name)[0]
            new_folder_name = f"{params['language']}_{base_name}"
            new_folder_path = os.path.join(DOWNLOAD_PATH, new_folder_name)
            os.makedirs(new_folder_path, exist_ok=True)

            docx_destination = os.path.join(new_folder_path, original_name)
            with open(docx_destination, "wb") as f:
                f.write(uploaded_file.read())

            paragraphs = load_text(docx_destination)
            model = load_model(params["model_name"], params["device"])

            st.write("Starting audio generation...")
            progress = st.progress(0)
            log_placeholder = st.empty()
            error_texts.clear()
            local_error_texts = []

            for index, paragraph in enumerate(paragraphs):
                output_file_name = generate_audio_filename(index)
                output_file_path = os.path.join(new_folder_path, output_file_name)

                st.write(f"Generating audio for paragraph {index + 1}/{len(paragraphs)}...")
                try:
                    log = tts_to_file_logged(
                        model,
                        text=paragraph,
                        out_path=output_file_path,
                        language=params["language"],
                        speaker_wav=params["voice"],
                        speed=0.85,
                    )
                    log_placeholder.code(log or "(no logs)")
                except Exception as e:
                    log_placeholder.error(f"Error generating audio: {e}")
                    if "exceeds the character limit" in str(e).lower():
                        local_error_texts.append(paragraph)
                    continue

                if "exceeds the character limit" in (log or "").lower():
                    new_name = f"audio_{index + 1}__may_have_error.wav"
                    os.rename(output_file_path, os.path.join(new_folder_path, new_name))
                    error_texts.append(paragraph)

                progress.progress((index + 1) / len(paragraphs))

            if local_error_texts or error_texts:
                error_docx_path = os.path.join(new_folder_path, "paragraphs_with_errors.docx")
                error_doc = Document()
                for t in local_error_texts + error_texts:
                    error_doc.add_paragraph(t)
                error_doc.save(error_docx_path)
                st.warning(f"Paragraphs with possible errors saved in: {error_docx_path}")

            st.success("Generation process completed!")
        finally:
            set_app_status(False)
            app_status_placeholder.markdown(
                "<div style='background-color:#4CAF50;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>✅ Available! No one is using it right now.</div>",
                unsafe_allow_html=True
            )

    if uploaded_file is None and not is_app_in_use():
        _console_standby_throttled("Waiting for user action.")

if __name__ == "__main__":
    main_app()
